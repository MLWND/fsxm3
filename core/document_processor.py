"""文档处理：加载 TXT/PDF/DOCX/CSV/Markdown → 分块 → 带 metadata。"""

import csv
import uuid
from datetime import datetime
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from loguru import logger

from config.settings import settings


def _read_text_file(file_path: str) -> str:
    """读取文本文件，自动检测编码（UTF-8 → GBK → latin-1）。"""
    raw = Path(file_path).read_bytes()
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _load_txt(file_path: str) -> list[Document]:
    text = _read_text_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_pdf(file_path: str) -> list[Document]:
    from langchain_community.document_loaders import PyPDFLoader
    return PyPDFLoader(file_path).load()


def _load_docx(file_path: str) -> list[Document]:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    text = "\n".join(
        para.text for para in doc.paragraphs if para.text.strip()
    )
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_csv(file_path: str) -> list[Document]:
    """CSV：每行作为一个 Document，内容为 "列名: 值" 格式。"""
    text = _read_text_file(file_path)
    reader = csv.reader(text.splitlines())
    rows = list(reader)
    if not rows:
        return []

    headers = rows[0] if rows else []
    docs = []
    for i, row in enumerate(rows[1:], start=1):
        # 将每行格式化为 "列1: 值1\n列2: 值2\n..."
        parts = []
        for h, v in zip(headers, row, strict=False):
            if v.strip():
                parts.append(f"{h}: {v.strip()}")
        if parts:
            content = "\n".join(parts)
            docs.append(Document(
                page_content=content,
                metadata={"source": file_path, "row": i},
            ))
    return docs if docs else [Document(page_content=text, metadata={"source": file_path})]


def _load_markdown(file_path: str) -> list[Document]:
    """Markdown：保留原始内容，直接加载为文本。"""
    text = _read_text_file(file_path)
    return [Document(page_content=text, metadata={"source": file_path})]


def _load_xlsx(file_path: str) -> list[Document]:
    """Excel：每个 sheet 的每行作为一个 Document，内容为 "列名: 值" 格式。"""
    from openpyxl import load_workbook

    docs = []
    wb = load_workbook(file_path, read_only=True, data_only=True)
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h) if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
        for row_idx, row in enumerate(rows[1:], start=1):
            parts = []
            for h, v in zip(headers, row, strict=False):
                if v is not None and str(v).strip():
                    parts.append(f"{h}: {v}")
            if parts:
                docs.append(Document(
                    page_content="\n".join(parts),
                    metadata={"source": file_path, "sheet": sheet_name, "row": row_idx},
                ))
    wb.close()
    return docs if docs else [Document(page_content="(空 Excel 文件)", metadata={"source": file_path})]


LOADERS = {
    ".txt": _load_txt,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
    ".csv": _load_csv,
    ".md": _load_markdown,
    ".markdown": _load_markdown,
    ".xlsx": _load_xlsx,
}


def get_text_splitter() -> RecursiveCharacterTextSplitter:
    return RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        add_start_index=True,
        separators=["\n\n", "\n", "。", "！", "？", "；", ".", "!", "?", ";", " "],
    )


def process_document(
    file_path: str,
    document_id: str | None = None,
) -> tuple[list[Document], str]:
    """加载文档、分块、附加 metadata。

    Returns:
        (chunks, document_id)
    """
    file_path = str(Path(file_path).resolve())
    suffix = Path(file_path).suffix.lower()

    if suffix not in LOADERS:
        raise ValueError(f"不支持的文件格式: {suffix}")

    document_id = document_id or str(uuid.uuid4())
    filename = Path(file_path).name

    logger.info("加载文档: {} (format={})", filename, suffix)
    raw_docs = LOADERS[suffix](file_path)
    logger.info("原始文档: {} 页/段", len(raw_docs))

    splitter = get_text_splitter()
    chunks = splitter.split_documents(raw_docs)
    logger.info("分块完成: {} chunks", len(chunks))

    now = datetime.now().isoformat()
    for i, chunk in enumerate(chunks):
        page = chunk.metadata.get("page", chunk.metadata.get("page_number", -1))
        chunk.metadata.update({
            "document_id": document_id,
            "source": filename,
            "page": page,
            "chunk_index": i,
            "total_chunks": len(chunks),
            "upload_time": now,
        })

    return chunks, document_id
